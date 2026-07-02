"""
file_parser.py — Utility module for extracting text from uploaded PDF and DOCX files.

This module provides functions to extract plain text from common document formats.
All extracted text is cleaned (excessive whitespace/newlines stripped) before being
returned so that the downstream AI agent receives clean input.

Supported formats: .pdf, .docx, .txt
"""

import io
import re
import logging
from typing import Optional

logger = logging.getLogger(__name__)


def _clean_text(text: str) -> str:
    """
    Normalize whitespace in extracted text:
    - Collapse multiple blank lines into a single blank line
    - Strip leading/trailing whitespace from each line
    - Strip leading/trailing whitespace from the entire result
    """
    # Replace multiple consecutive newlines (with optional whitespace between) with double newline
    text = re.sub(r'\n\s*\n', '\n\n', text)
    # Collapse runs of spaces/tabs (but not newlines) into a single space
    text = re.sub(r'[^\S\n]+', ' ', text)
    # Strip each line individually, then rejoin
    lines = [line.strip() for line in text.splitlines()]
    text = '\n'.join(lines)
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file given as raw bytes.

    Uses pdfplumber for high-quality text extraction, including PDFs with
    complex layouts, tables, and multi-column text.

    Raises:
        ValueError: If the PDF contains no extractable text (e.g. scanned/image-based PDF).
        Exception: If the PDF is malformed or cannot be opened.
    """
    import pdfplumber

    pages_text = []

    try:
        # pdfplumber can open from a BytesIO stream
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            logger.info(f"PDF has {len(pdf.pages)} page(s)")

            for i, page in enumerate(pdf.pages):
                # extract_text() returns None if the page has no text layer
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
                else:
                    logger.warning(f"PDF page {i + 1} has no extractable text (may be an image)")

    except Exception as e:
        logger.error(f"Failed to open/parse PDF: {e}")
        raise ValueError(f"Failed to parse PDF file: {str(e)}")

    # If no pages had any text, it's likely a scanned/image-based PDF
    full_text = "\n\n".join(pages_text)
    if not full_text.strip():
        raise ValueError(
            "No extractable text found in PDF — it may be a scanned image. "
            "Please use a PDF with selectable text, or paste the text manually."
        )

    return _clean_text(full_text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX file given as raw bytes.

    Extracts text from:
    1. All paragraphs in the document body
    2. All table cells (to capture tabular resume data)

    Raises:
        ValueError: If the DOCX is empty or cannot be parsed.
    """
    from docx import Document

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as e:
        logger.error(f"Failed to open/parse DOCX: {e}")
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")

    parts = []

    # Extract text from all paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract text from all table cells — resumes and JDs often use tables for layout
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_texts.append(cell_text)
            if row_texts:
                parts.append(" | ".join(row_texts))

    full_text = "\n".join(parts)
    if not full_text.strip():
        raise ValueError(
            "No text content found in DOCX file. "
            "The document appears to be empty."
        )

    return _clean_text(full_text)


def extract_text_from_upload(filename: str, file_bytes: bytes) -> str:
    """
    Dispatcher function: detect the file type by extension and route to
    the appropriate extractor.

    Args:
        filename: Original filename (used to detect extension).
        file_bytes: Raw bytes of the uploaded file.

    Returns:
        Extracted plain text from the file.

    Raises:
        ValueError: If the file type is unsupported or parsing fails.
    """
    # Normalize the extension to lowercase for case-insensitive matching
    extension = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    logger.info(f"Extracting text from file: {filename} (extension: .{extension})")

    if extension == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif extension == 'docx':
        return extract_text_from_docx(file_bytes)
    elif extension == 'txt':
        # Plain text files: just decode and clean
        try:
            text = file_bytes.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 which never raises (covers most legacy text files)
            text = file_bytes.decode('latin-1')
        if not text.strip():
            raise ValueError("The uploaded .txt file is empty.")
        return _clean_text(text)
    else:
        raise ValueError(
            f"Unsupported file type: .{extension}. "
            f"Please upload a .pdf, .docx, or .txt file."
        )
